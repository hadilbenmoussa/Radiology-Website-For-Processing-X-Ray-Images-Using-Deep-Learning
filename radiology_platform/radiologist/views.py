from django.urls import reverse_lazy
from django.http import HttpResponseRedirect
from django.shortcuts import render,redirect,get_object_or_404
import docx
import os
import io
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.shared import Pt,RGBColor

from PIL import Image
from django.http import FileResponse
from django.contrib import messages
from django.contrib.auth import logout
from django.db.models import Q
from .forms import PatientForm,ReportForm,AppointmentForm,ReportFileForm
from doctor.models import Patient,User,Report
from .models import Appointment
from .decorators import radiologist_required
from django.contrib.auth.decorators import login_required
from django.http import HttpResponse
from django.db.models import Case, When,Value,IntegerField
from django.core.files.storage import FileSystemStorage
import keras
from keras.models import load_model
import keras.utils as image
import numpy as np
import tensorflow as tf
import json
from tensorflow import Graph



# Create your views here.
@radiologist_required
def radiologistHomeView(request):
    return render (request,'radiologist/dashboard.html')


@radiologist_required
def appointmentsView(request):
    appointments = Appointment.objects.filter(radiologist=request.user)

  
    return render(request, 'radiologist/appointments.html', {'appointments': appointments})


@radiologist_required
def delete_report(request, report_id):
    report= Report.objects.get(id=report_id)
    report.delete()
    return redirect('rad_reports')

@radiologist_required
def delete_appointment(request, appointment_id):
    appointment = Appointment.objects.get(id=appointment_id)
    
    appointment.delete()
    return redirect('rad_appointments')

@radiologist_required    
def medical_history(request, patient_id):
    patient = get_object_or_404(Patient, id=patient_id)
    reports = Report.objects.filter(patient=patient)
    context={'patient': patient, 'reports': reports}
    return render(request, 'radiologist/medical_history.html', context)


@radiologist_required    
def reportsView(request):
    reports = Report.objects.filter(radiologist=request.user).order_by(Case(
        When(urgency='U', then=Value(1)),
        When(urgency='N', then=Value(2)),
        When(urgency='L', then=Value(3)),
        default=Value(4),
        output_field=IntegerField(),
    )) 

    context = {
        'reports': reports,
    }    
    return render(request, 'radiologist/reports.html', context)




#ken tla3lek erreur 3al image dimensions baddelhom hne
img_height, img_width=180,180

#
model = keras.models.load_model('./models/model.h5')

@radiologist_required  
def generate_report(request, report_id):
    prediction={}
    report = get_object_or_404(Report, id=report_id)
    date='11/12/2023'
    print(date)
    if request.method == 'POST':
        form = ReportFileForm(request.POST, request.FILES)
        if form.is_valid():

            # get form data
            image = form.cleaned_data['image']
            findings = form.cleaned_data['findings']
            impression = form.cleaned_data['impression']
            recommendations = form.cleaned_data['recommendations']
            indications = form.cleaned_data['indications']
            AI_check = request.POST.get('AI_check')

                           # get the uploaded file from the POST request
            

            # create new Word document
            document = docx.Document()
            
            # Add section properties to the document and set border properties
            sect_pr = document.element.body.sectPr
            pg_borders = docx.oxml.shared.OxmlElement('w:pgBorders')
            pg_borders.set(docx.oxml.ns.qn('w:top'), 'single 1px #000000')
            pg_borders.set(docx.oxml.ns.qn('w:right'), 'single 1px #000000')
            pg_borders.set(docx.oxml.ns.qn('w:bottom'), 'single 1px #000000')
            pg_borders.set(docx.oxml.ns.qn('w:left'), 'single 1px #000000')
            sect_pr.append(pg_borders)
            section = document.sections[0]
            section.left_margin = docx.shared.Inches(1.5)
            section.right_margin = docx.shared.Inches(1)
            header = section.header
            table = header.add_table(rows=2, cols=3, width=docx.shared.Inches(8))
            table.style = 'Table Grid'
            patient_info_cell = table.cell(0, 0)
            patient_info_cell.text = 'Patient Info'
            patient_info_cell.paragraphs[0].runs[0].bold = True
            patient_info = 'Name: {}\nAge: {}\nGender: {}'.format(report.patient.get_patient_name(),20,'Female')
            patient_info_cell = table.cell(1, 0)
            patient_info_cell.text = patient_info
            patient_info_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            report_id_cell = table.cell(0, 1)
            report_id_cell.text = 'Report Data'
            report_id_cell.paragraphs[0].runs[0].bold = True
            report_id_cell = table.cell(1, 1)
            report_id_info = 'Exam ID: {}\nDate&Time: {}\nOrdering Provider Dr: {}'.format(report.id, date, report.doctor)
            report_id_cell.text = report_id_info
            report_id_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            addition_info_cell = table.cell(0, 2)
            addition_info_cell.text = 'Additional Data'
            addition_info_cell.paragraphs[0].runs[0].bold = True
            addition_info = 'Name: {}\nAge: {}\nGender: {}'.format(report.exam_type,report.body_part,report.urgency)
            addition_info_cell = table.cell(1, 2)
            addition_info_cell.text = addition_info
            addition_info_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT



            blank_space = document.add_paragraph()
            blank_space.space_before = docx.shared.Pt(12)

            print('fine')
            # add image to the document
            img_data = io.BytesIO(image.read())
            if image is not None:
            # save the file to the media directory
                  fs = FileSystemStorage()
                  filePathName = fs.save(image.name, image)
                  filePathName = fs.url(filePathName)
                  print(filePathName)
                  testimage = '.' + filePathName
                  print(testimage)
            # load and preprocess the image using tensorflow
                  img = tf.keras.utils.load_img(testimage, target_size=(img_height, img_width))
                  x = tf.keras.utils.img_to_array(img)
                  x = np.expand_dims(x, axis=0)
            # perform model inference
                  predictions = model.predict(x)
                  score = tf.nn.softmax(predictions[0])
                  predictedLabel = ["NORMAL", "PNEUMONIA"]
            # prepare the context for rendering the response template
                  prediction = {'filePathName': filePathName, 'predictedLabel': predictedLabel[np.argmax(score)], 'percentage': 100 * np.max(score)}      
                  aidiagnosis='The classification of the scan is "{}" with "{}" % confidence'.format(predictedLabel[np.argmax(score)],100 * np.max(score))

            centered = document.add_paragraph()
            centered.alignment = WD_ALIGN_PARAGRAPH.CENTER
            styles = document.styles

            # add the image to the centered paragraph
            heading_style = styles['Heading 1']
            heading_style.font.name = 'Arial'
            heading_style.font.size = Pt(14)
            heading_style.font.bold = True
            heading_style.font.color.rgb = RGBColor(0, 0, 0)


            text_style = styles['Normal']
            text_style.font.name = 'Arial'
            text_style.font.size = Pt(10)
            text_style.font.color.rgb = RGBColor(0, 0, 0)
            run = centered.add_run()
            run.add_picture(img_data, width=docx.shared.Inches(6), height=docx.shared.Inches(4))
            print('not fine')
            document.add_heading('Indications', level=1)
            document.add_paragraph(indications, style=text_style)

            document.add_heading('Findings', level=1)
            document.add_paragraph(findings, style=text_style)

            document.add_heading('Impression', level=1)
            document.add_paragraph(impression, style=text_style)
            
            document.add_heading('Recommendations', level=1)
            document.add_paragraph(recommendations, style=text_style)
            if AI_check:
                            document.add_heading('AI Pneumonia Diagnosis', level=1)
                            document.add_paragraph(aidiagnosis, style=text_style)
            directory_path1 = 'reports'
            BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

            directory_path = os.path.join(BASE_DIR, 'reports')
            if not os.path.exists(directory_path):
                os.makedirs(directory_path)
            file_name = 'report{}.docx'.format(report.id)
            file_path = os.path.join(directory_path, file_name)
            document.save(file_path)
            file_path1 = os.path.join(directory_path1, file_name)
            document.save(file_path1)
            report.report_file = file_path
            report.status='P'
            report.appointment_set.all().delete()
            report.save()
            

            # return the document as a downloadable response
            with open(file_path, 'rb') as f:
                response = HttpResponse(f, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                response['Content-Disposition'] = f'attachment; filename={file_name}'

            return response
        else:
            print(form.errors)

    else:
        print('not post request')
        form = ReportFileForm()
    return render(request, 'radiologist/generate_report.html', {'report': report, 'form': form,'prediction':prediction})



@radiologist_required  
def create_report(request):
    if request.method == 'POST':
        form = ReportForm(request.POST)
        if form.is_valid():
            report = form.save(commit=False)
            report.radiologist = request.user
            report.save()
    else:
        form = ReportForm()
    if 'HX-Request' in request.headers:    
        base_template = 'radiologist/create_report.html'
    else :
        base_template ='radiologist/reports.html'
    return render(request, base_template, {'form': form})

  
def report_details(request, report_id):
    report = get_object_or_404(Report, id=report_id)
    if not('HX-Request' in request.headers):    
        return render(request, 'radiologist/reports.html')
    else:
        return render(request, 'radiologist/report_details.html',{'report':report})
    
    

@radiologist_required    
def create_appointment(request, report_id):
    report = get_object_or_404(Report, id=report_id)
    initial_values = {
    'email': report.patient.email,
    'exam_type': report.exam_type,
    'details': report.details,}
    form = AppointmentForm(initial=initial_values)
    if request.method == 'POST':
        form = AppointmentForm(request.POST, initial=initial_values)
        if form.is_valid():
            appointment = form.save(commit=False)
            appointment.radiologist = report.radiologist
            appointment.patient = report.patient
            appointment.report = report
            appointment.body_part = report.body_part
            appointment.status = 'S'
            appointment.save()

    context = {
        'form': form,
        'report': report,
    }
    return render(request, 'radiologist/create_appointment.html', context)

@radiologist_required    
def update_appointment(request,appointment_id):
    if request.method == 'POST':
        appointment=Appointment.objects.get(id=appointment_id)
        report=appointment.report    
        form=AppointmentForm(request.POST,instance=appointment)
        if form.is_valid():
            form.save()
            print('valid')
            return HttpResponseRedirect(reverse_lazy('rad_appointments'))
        else :
            print(form.errors)    
    else:
        
        appointment=Appointment.objects.get(id=appointment_id)
        report=appointment.report   
        form=AppointmentForm(instance=appointment) 
    return render(request,'radiologist/update_appointment.html',{'form':form})


    


@radiologist_required    
def blogView(request):
    return render (request,'radiologist/addblog.html')
@radiologist_required    
def profileView(request,username):
    return render (request,'radiologist/profile.html')
@radiologist_required    
def notificationsView(request):
    return render (request,'radiologist/notifications.html')
@radiologist_required
def create_patient(request):
    if request.method == 'POST':
        form = PatientForm(request.POST)
        print('form received')
        if form.is_valid():
            email = form.cleaned_data['email']
            cin_number = form.cleaned_data['cin_number']
            existing_patient = get_object_or_none(Patient, email=email,cin_number=cin_number)
            if existing_patient:
                messages.success(request,'Already registered')
            else:
                messages.success(request,'registration successful')
                form.save()

            return redirect('patients')
    else:
        form = PatientForm()
    return render(request, 'radiologist/create_patient_form.html', {'form': form})

@login_required
def logout_user(request):
    logout(request)
    return redirect('login_view')
 
def get_object_or_none(model, email=None, cin_number=None, **kwargs):
    filter_condition = Q()
    if email is not None:
        filter_condition |= Q(email=email)
    if cin_number is not None:
        filter_condition |= Q(cin_number=cin_number)
    if kwargs:
        filter_condition &= Q(**kwargs)

    try:
        return model.objects.get(filter_condition)
    except model.DoesNotExist:
        return None    









# Create your views here.
def patientHomeView(request):
    return render (request,'patient/home.html')
def index(request):
    context={'a':1}
    return render(request,'index.html',context)

@radiologist_required
def predictImage(request):
      if request.method == 'GET':
           return render(request,'radiologist/index.html')
      elif request.method == 'POST':
        # get the uploaded file from the POST request
        fileObj = request.FILES.get('file')
        print(fileObj)
        if fileObj is not None:
            # save the file to the media directory
            fs = FileSystemStorage()
            filePathName = fs.save(fileObj.name, fileObj)
            filePathName = fs.url(filePathName)
            print(filePathName)
            testimage = '.' + filePathName
            print(testimage)
            # load and preprocess the image using tensorflow
            img = tf.keras.utils.load_img(testimage, target_size=(img_height, img_width))
            x = tf.keras.utils.img_to_array(img)
            x = np.expand_dims(x, axis=0)
            # perform model inference
            predictions = model.predict(x)
            score = tf.nn.softmax(predictions[0])
            predictedLabel = ["NORMAL", "PNEUMONIA"]
            # prepare the context for rendering the response template
            context = {'filePathName': filePathName, 'predictedLabel': predictedLabel[np.argmax(score)], 'percentage': 100 * np.max(score)}
            # render the response template with the context
            return render(request, 'radiologist/index.html', context)
  
  
def viewDataBase(request):
    import os
    listOfImages=os.listdir('./media/')
    listOfImagesPath=['./media/'+i for i in listOfImages]
    context={'listOfImagesPath':listOfImagesPath}
    return render(request,'viewDB.html',context)     